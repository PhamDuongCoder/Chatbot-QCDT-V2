#!/usr/bin/env python3
"""
Unified Pipeline: Agentic Preprocessing + Embedding
File này gộp tất cả quy trình tiền xử lý, chunking, và embedding vào một file duy nhất.
Chỉ cần truyền vào đường dẫn file và run script này.
"""

import os
import sys
import time
import re
import json
import csv
import fitz
import tempfile
import tomllib
from pathlib import Path
from dotenv import load_dotenv
from google import genai
from google.genai import types
import psycopg2
from pgvector.psycopg2 import register_vector
from docx2pdf import convert

# Load environment variables
load_dotenv()

# ============================================================================
# LOAD CITATION MAPPING AT MODULE LEVEL
# ============================================================================

def load_citations():
    """Load Citation.csv as a dict {parent_doc_id: url}"""
    citation_path = Path(__file__).parent.parent / "Citation.csv"
    citation_dict = {}
    try:
        with open(citation_path, encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                citation_dict[row["parent_doc_id"].strip()] = row["url"].strip()
        print(f"[✓] Loaded {len(citation_dict)} citations from Citation.csv")
    except Exception as e:
        print(f"[!] Warning: Could not load Citation.csv: {e}")
    return citation_dict

CITATIONS = load_citations()

# ============================================================================
# CONFIGURATION
# ============================================================================

# Gemini and Google API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client(api_key=GOOGLE_API_KEY)
MODEL = "gemini-3.1-flash-lite"
EMBEDDING_MODEL = "gemini-embedding-001"

# Paths
DATA_DIR = Path(__file__).parent.parent / "Data"
PREPROCESSED_DIR = Path(__file__).parent.parent / "Preprocessed_Data"

# Parameters
PAGE_THRESHOLD = 8
WORDS_PER_PAGE = 400

# Database Configuration
with open("../.streamlit/secrets.toml", "rb") as f:
    secrets = tomllib.load(f)

DB_CONFIG = {
    "host": secrets.get("SUPABASE_DB_HOST", "localhost"),
    "port": secrets.get("SUPABASE_DB_PORT", "5432"),
    "database": secrets.get("SUPABASE_DB_NAME"),
    "user": secrets.get("SUPABASE_DB_USER"),
    "password": secrets.get("SUPABASE_DB_PASSWORD")
}
# ============================================================================
# SYSTEM INSTRUCTIONS
# ============================================================================

SYSTEM_INSTRUCTION = """
Bạn là chuyên gia tiền xử lý dữ liệu RAG cho hệ thống Quy chế đào tạo ĐHBK Hà Nội.

NHIỆM VỤ:
Phân tích tài liệu PDF hoặc .docx, sửa lỗi trình bày, chuyển bảng sang Markdown, và chia thành các chunk logic.
Giữ nguyên TOÀN BỘ nội dung tài liệu, bao gồm tất cả các điểm a), b), c)..., hệ số, mức phí, điều kiện cụ thể. Tuyệt đối không lược bỏ bất kỳ điều khoản hoặc chi tiết số liệu nào, dù nhỏ.

THÔNG TIN TÀI LIỆU:
- Tên file: {filename}
- parent_doc_id: {parent_doc_id}
- category: {category}
- year: {year}

CHIẾN LƯỢC CHIA CHUNK:
- Đơn vị cơ bản: mỗi Điều = 1 chunk
- Nếu Điều quá ngắn (< 80 từ): gộp với Điều liền kề cùng chủ đề
- Nếu Điều quá dài (> 600 từ): tách theo Khoản, mỗi Khoản là 1 chunk
- Bảng biểu lớn (> 5 hàng): có thể là chunk độc lập
- Phần mở đầu/định nghĩa chung: 1 chunk riêng
- Ranh giới Chương/Mục lớn (I, II, III... hoặc tương đương) LUÔN bắt đầu chunk mới, KHÔNG được gộp với nội dung của Chương/Mục trước dù đoạn đó < 80 từ

QUY TẮC ĐỊNH DẠNG:
1. Mỗi chunk nằm trong <<<CHUNK_START>>> ... <<<CHUNK_END>>>
2. Không trả về bất kỳ lời dẫn, giải thích hay hội thoại nào ngoài các thẻ trên
3. Bảng biểu dùng định dạng Markdown table
4. Hình ảnh/biểu đồ: mô tả bằng [Hình: <mô tả chi tiết tất cả những yếu tố quan trọng trong hình ảnh/biểu đồ có thể xuất hiện trong query>]
5. Trả lời hoàn toàn bằng tiếng Việt

CHIẾN LƯỢC METADATA CHO BẢNG BIỂU:
- topic_tags: liệt kê tất cả tên chương trình, học phần, đối tượng, mã số được đề cập trong bảng
  (ví dụ: nếu bảng có Global ICT, Logistics, Việt-Nhật thì tất cả phải có trong topic_tags)
- summary: phải đề cập ít nhất 2-3 đối tượng/giá trị cụ thể từ bảng, không chỉ mô tả chung chung
  Sai:  "Bảng học phí các chương trình đặc biệt ELITECH"
  Đúng: "Bảng học phí theo TCHP cho các chương trình ELITECH gồm Global ICT, Logistics,
         Việt-Nhật, với mức từ X đến Y triệu đồng/TCHP"

CẤU TRÚC CHUNK (tuân thủ chính xác, không thêm/bớt field):
<<<CHUNK_START>>>
source: {filename}
parent_doc_id: {parent_doc_id}
chunk_id: {doc_id}_001  ← tăng dần, 3 chữ số, ví dụ _002, _003
chunk_index: 1          ← số nguyên tăng dần
language: vi
category: {category}
year: {year}
chunk_title: [Tên điều khoản hoặc chủ đề chính]
topic_tags: [tag1, tag2, tag3]
summary: [2-3 câu tóm tắt, ưu tiên từ khóa: mã học phần, tín chỉ, cảnh báo học tập, điều kiện xét]

[Nội dung chunk ở đây]
<<<CHUNK_END>>>
"""

STRUCTURE_PROMPT = """
Phân tích cấu trúc tài liệu này và trả về JSON theo format sau. KHÔNG trả về gì khác ngoài JSON thuần:
[
  {"section": "Chương I - Quy định chung", "start_page": 1, "end_page": 4},
  {"section": "Chương II - ...", "start_page": 5, "end_page": 11}
]
Quy tắc:
- Đơn vị chia là Chương hoặc phần lớn tương đương (nếu không có Chương thì chia theo nhóm chủ đề)
- Mỗi phần tối thiểu 5 trang, tối đa 10 trang
- Ưu tiên tạo ÍT phần nhất có thể, gộp các Chương ngắn liền kề nếu tổng không vượt 10 trang
- Nếu 1 Chương > 10 trang: tách theo nhóm Điều, nhưng KHÔNG BAO GIỜ cắt giữa một Điều
- Không có markdown, không có code block, chỉ JSON
"""

# ============================================================================
# DATABASE FUNCTIONS
# ============================================================================

def get_connection():
    """Get a database connection"""
    conn = psycopg2.connect(**DB_CONFIG)
    register_vector(conn)
    return conn

def setup_database():
    """Create database schema if it doesn't exist"""
    create_table_sql = '''
    CREATE TABLE IF NOT EXISTS chunks (
        id SERIAL PRIMARY KEY,
        chunk_id TEXT UNIQUE NOT NULL,
        parent_doc_id TEXT NOT NULL,
        source TEXT,
        category TEXT,
        year TEXT,
        chunk_index INTEGER,
        chunk_title TEXT,
        topic_tags TEXT,
        summary TEXT,
        content TEXT NOT NULL,
        url TEXT DEFAULT '',
        embedding vector(3072)
    );
    '''
    add_url_column_sql = '''
    ALTER TABLE chunks ADD COLUMN IF NOT EXISTS url TEXT DEFAULT '';
    '''
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(create_table_sql)
                cur.execute(add_url_column_sql)
            conn.commit()
        print("[✓] Database schema ready")
    except Exception as e:
        print(f"[!] Database error: {e}")

# ============================================================================
# PREPROCESSING FUNCTIONS
# ============================================================================

def convert_docx_to_temp_pdf(docx_path: Path) -> Path:
    """Convert .docx sang PDF tạm, trả về path của PDF"""
    tmp_dir = Path(tempfile.mkdtemp())
    pdf_path = tmp_dir / docx_path.with_suffix(".pdf").name
    convert(docx_path, pdf_path)
    return pdf_path

def get_page_count(file_path: Path) -> int:
    """Đếm trang PDF hoặc ước tính số trang cho docx"""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        doc = fitz.open(file_path)
        count = doc.page_count
        doc.close()
        return count
    elif suffix == ".docx":
        from docx import Document
        doc = Document(file_path)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        return max(1, word_count // WORDS_PER_PAGE)
    return 0

def extract_structure_from_pdf(file_path: Path) -> list[dict]:
    """Pass 1: Gemini phân tích cấu trúc PDF, trả về list các section với page range"""
    print("    [~] Phân tích cấu trúc tài liệu...")
    file_upload = client.files.upload(file=file_path)

    while file_upload.state.name == "PROCESSING":
        time.sleep(2)
        file_upload = client.files.get(name=file_upload.name)

    if file_upload.state.name == "FAILED":
        client.files.delete(name=file_upload.name)
        raise RuntimeError("Upload thất bại khi phân tích cấu trúc")

    response = client.models.generate_content(
        model=MODEL,
        contents=[file_upload, STRUCTURE_PROMPT],
        config=types.GenerateContentConfig(temperature=0.0)
    )
    client.files.delete(name=file_upload.name)

    # Xóa markdown code fence nếu Gemini vẫn thêm vào
    text = re.sub(r'^```(?:json)?\s*|\s*```$', '', response.text.strip())
    return json.loads(text)

def split_pdf(file_path: Path, structure: list[dict], output_folder: Path) -> list[Path]:
    """Cắt PDF theo cấu trúc từ Gemini"""
    doc = fitz.open(file_path)
    parts = []
    for i, section in enumerate(structure):
        start = section["start_page"] - 1   # fitz dùng 0-index
        end   = section["end_page"] - 1
        sub   = fitz.open()
        sub.insert_pdf(doc, from_page=start, to_page=end)
        part_path = output_folder / f"{file_path.stem}_{i+1:02d}.pdf"
        sub.save(part_path)
        sub.close()
        parts.append(part_path)
    doc.close()
    return parts

def split_docx(file_path: Path, output_folder: Path) -> list[Path]:
    """Cắt docx theo ranh giới Chương, đảm bảo mỗi part ≤ PAGE_THRESHOLD trang ước tính"""
    import copy
    from docx import Document

    doc = Document(file_path)
    chapter_re = re.compile(r'^(Chương|CHƯƠNG)\s+[IVXLCDM\d]+', re.IGNORECASE)

    sections, current, current_words = [], [], 0
    for para in doc.paragraphs:
        is_boundary = chapter_re.match(para.text.strip())
        words = len(para.text.split())

        if is_boundary and current and current_words > WORDS_PER_PAGE * 2:
            sections.append(current)
            current, current_words = [], 0

        current.append(para)
        current_words += words

        if current_words >= WORDS_PER_PAGE * PAGE_THRESHOLD:
            sections.append(current)
            current, current_words = [], 0

    if current:
        sections.append(current)

    parts = []
    for i, paras in enumerate(sections):
        new_doc = Document()
        # Xóa paragraph mặc định rỗng
        for p in new_doc.paragraphs:
            p._element.getparent().remove(p._element)
        for para in paras:
            new_doc.add_paragraph()._p.getparent().replace(
                new_doc.paragraphs[-1]._p,
                copy.deepcopy(para._p)
            )
        part_path = output_folder / f"{file_path.stem}_{i+1:02d}.docx"
        new_doc.save(part_path)
        parts.append(part_path)
    return parts

def upload_and_process(file_path: Path, original_doc_id: str = None, original_category: str = None) -> Path:
    """Upload file to Gemini, preprocess with chunking, save to Preprocessed_Data folder"""
    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
    if not file_path.exists() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        print(f"[!] Bỏ qua (không hỗ trợ): {file_path.name}")
        return None

    category = file_path.parent.name
    filename = file_path.name
    doc_id = file_path.stem
    parent_doc_id = original_doc_id if original_doc_id else doc_id

    if original_doc_id and original_category:
        base_folder = PREPROCESSED_DIR / original_category / f"{original_doc_id}_partitioned_processed"
    else:
        base_folder = PREPROCESSED_DIR / category

    year_match = re.search(r'\b(19|20)\d{2}\b', doc_id)
    year = year_match.group() if year_match else "N/A"

    print(f"\n[*] Tiền xử lý: {filename} (Category: {category}, Year: {year})")

    temp_pdf = None
    upload_path = file_path
    if file_path.suffix.lower() == ".docx":
        print(f"    [~] Chuyển đổi .docx → .pdf...")
        temp_pdf = convert_docx_to_temp_pdf(file_path)
        upload_path = temp_pdf

    try:
        file_upload = client.files.upload(file=upload_path)
        while file_upload.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(2)
            file_upload = client.files.get(name=file_upload.name)

        if file_upload.state.name == "FAILED":
            print(f"\n[X] Upload thất bại: {filename}")
            return None

        formatted_instruction = SYSTEM_INSTRUCTION.format(
            filename=filename,
            doc_id=doc_id,
            parent_doc_id=parent_doc_id,
            category=original_category or category,
            year=year
        )

        response = client.models.generate_content(
            model=MODEL,
            contents=[file_upload, "Hãy tiền xử lý tài liệu này theo đúng hướng dẫn."],
            config=types.GenerateContentConfig(
                system_instruction=formatted_instruction,
                temperature=0.1
            )
        )

        if "<<<CHUNK_START>>>" not in response.text:
            print(f"\n[!] Output không đúng format — lưu vào thư mục review")
            output_folder = PREPROCESSED_DIR / "_review" / (original_category or category)
        else:
            output_folder = base_folder

        output_folder.mkdir(parents=True, exist_ok=True)
        output_file = output_folder / f"{doc_id}_preprocessed.txt"

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(response.text.strip())

        client.files.delete(name=file_upload.name)
        print(f"\n[✓] Xong: {output_file}")
        return output_file

    except Exception as e:
        print(f"\n[X] Lỗi khi xử lý {filename}: {e}")
        return None

    finally:
        if temp_pdf and temp_pdf.exists():
            temp_pdf.unlink()
            temp_pdf.parent.rmdir()

def partition_and_process(file_path: Path) -> list[Path]:
    """Xử lý file lớn: partition nếu chưa có, rồi gửi từng part"""
    original_doc_id = file_path.stem
    original_category = file_path.parent.name
    partition_folder = file_path.parent / f"{original_doc_id}_partitioned"
    marker = partition_folder / "partitioned"

    if marker.exists():
        print("    [~] Tìm thấy partition cũ, dùng lại...")
        parts = sorted(f for f in partition_folder.iterdir()
                       if f.suffix.lower() in {".pdf", ".docx"})
    else:
        partition_folder.mkdir(parents=True, exist_ok=True)
        if file_path.suffix.lower() == ".pdf":
            structure = extract_structure_from_pdf(file_path)
            parts = split_pdf(file_path, structure, partition_folder)
        else:
            parts = split_docx(file_path, partition_folder)
        marker.touch()
        print(f"    [~] Đã tạo {len(parts)} parts tại {partition_folder.name}/")

    preprocessed_files = []
    for part in parts:
        output_file = upload_and_process(part, original_doc_id=original_doc_id, original_category=original_category)
        if output_file:
            preprocessed_files.append(output_file)
    
    return preprocessed_files

def preprocess_file(file_path: Path) -> list[Path]:
    """Entry point cho preprocessing: kiểm tra kích thước, partition nếu cần"""
    file_path = Path(file_path)
    if not file_path.exists() or file_path.suffix.lower() not in {".pdf", ".docx"}:
        print(f"[!] Bỏ qua (không hỗ trợ): {file_path.name}")
        return []

    page_count = get_page_count(file_path)
    suffix_note = "(ước tính)" if file_path.suffix.lower() == ".docx" else ""
    print(f"\n[*] {file_path.name} — {page_count} trang {suffix_note}")

    if page_count <= PAGE_THRESHOLD:
        output_file = upload_and_process(file_path)
        return [output_file] if output_file else []
    else:
        print(f"    [!] Vượt ngưỡng {PAGE_THRESHOLD} trang → partition mode")
        return partition_and_process(file_path)

# ============================================================================
# EMBEDDING FUNCTIONS
# ============================================================================

def parse_preprocessed_file(file_path: Path) -> list[dict]:
    """Parse preprocessed file to extract chunks"""
    chunks = []
    with open(file_path, encoding="utf-8") as f:
        content = f.read()
    chunk_blocks = re.split(r"<<<CHUNK_START>>>|<<<CHUNK_END>>>", content)
    for block in chunk_blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        meta = {}
        i = 0
        while i < len(lines) and lines[i].strip():
            if ':' in lines[i]:
                k, v = lines[i].split(':', 1)
                meta[k.strip()] = v.strip()
            i += 1
        while i < len(lines) and not lines[i].strip():
            i += 1
        content_text = '\n'.join(lines[i:]).strip()
        meta['content'] = content_text
        if meta.get('chunk_id') and meta.get('content'):
            chunks.append(meta)
    return chunks

def embed_chunk(chunk, max_retries=3):
    """Embed a chunk using Gemini embedding model"""
    text = (chunk.get('summary', '') + "\n\n" + chunk.get('content', '')).strip()
    for attempt in range(max_retries):
        try:
            response = client.models.embed_content(
                model=EMBEDDING_MODEL,
                contents=text,
                config=types.EmbedContentConfig(task_type="RETRIEVAL_DOCUMENT")
            )
            return response.embeddings[0].values
        except Exception as e:
            msg = str(e)
            if any(code in msg for code in ["429", "503"]):
                print(f"    [~] Rate limit, retry {attempt+1}/3...")
                time.sleep(30)
            else:
                print(f"    [!] Embedding error: {e}")
                break
    return None

def store_chunk(cur, chunk, embedding):
    """Store chunk in database"""
    sql = '''
    INSERT INTO chunks (chunk_id, parent_doc_id, source, category, year,
                        chunk_index, chunk_title, topic_tags, summary, content, url, embedding)
    VALUES (%(chunk_id)s, %(parent_doc_id)s, %(source)s, %(category)s, %(year)s,
            %(chunk_index)s, %(chunk_title)s, %(topic_tags)s, %(summary)s, %(content)s, %(url)s, %(embedding)s)
    ON CONFLICT (chunk_id) DO NOTHING;
    '''
    
    # Look up URL from CITATIONS dict
    parent_doc_id = chunk.get('parent_doc_id', '')
    url = CITATIONS.get(parent_doc_id, '')
    
    params = {
        'chunk_id': chunk.get('chunk_id'),
        'parent_doc_id': parent_doc_id,
        'source': chunk.get('source'),
        'category': chunk.get('category'),
        'year': chunk.get('year'),
        'chunk_index': int(chunk.get('chunk_index')) if chunk.get('chunk_index') else None,
        'chunk_title': chunk.get('chunk_title'),
        'topic_tags': chunk.get('topic_tags'),
        'summary': chunk.get('summary'),
        'content': chunk.get('content'),
        'url': url,
        'embedding': embedding
    }
    cur.execute(sql, params)

def embed_and_store_file(file_path: Path):
    """Parse preprocessed file, embed chunks, and store in database"""
    chunks = parse_preprocessed_file(file_path)
    print(f"\n[*] Embedding {len(chunks)} chunks từ {file_path.name}...")
    
    with get_connection() as conn:
        with conn.cursor() as cur:
            for i, chunk in enumerate(chunks, 1):
                try:
                    cur.execute("SELECT 1 FROM chunks WHERE chunk_id=%s", (chunk.get('chunk_id'),))
                    if cur.fetchone():
                        print(f"    [SKIP] {chunk.get('chunk_id')} (đã tồn tại)")
                        continue
                    
                    embedding = embed_chunk(chunk)
                    if embedding is None:
                        print(f"    [FAIL] {chunk.get('chunk_id')} (embedding error)")
                        continue
                    
                    store_chunk(cur, chunk, embedding)
                    conn.commit()
                    print(f"    [OK {i}/{len(chunks)}] {chunk.get('chunk_id')}")
                    time.sleep(0.5)
                    
                except Exception as e:
                    conn.rollback()
                    print(f"    [FAIL] {chunk.get('chunk_id')} ({e})")

# ============================================================================
# MAIN UNIFIED PIPELINE
# ============================================================================

def unified_pipeline(file_path_str: str):
    """
    Unified pipeline: Preprocessing + Embedding
    Takes a file path, preprocesses it, then embeds all chunks to database
    """
    print("=" * 70)
    print("UNIFIED PIPELINE: TIỀN XỬ LÝ + EMBEDDING")
    print("=" * 70)
    
    file_path = Path(file_path_str)
    
    # Step 1: Setup database
    print("\n[STEP 1] Khởi tạo database...")
    setup_database()
    
    # Step 2: Preprocessing
    print("\n[STEP 2] Tiền xử lý tài liệu...")
    preprocessed_files = preprocess_file(file_path)
    
    if not preprocessed_files:
        print("[X] Tiền xử lý thất bại. Dừng pipeline.")
        return
    
    print(f"\n[✓] Tạo được {len(preprocessed_files)} file(s) đã tiền xử lý")
    
    # Step 3: Embedding and storing
    print("\n[STEP 3] Embedding và lưu vào database...")
    for preprocessed_file in preprocessed_files:
        if preprocessed_file.exists():
            embed_and_store_file(preprocessed_file)
    
    print("\n" + "=" * 70)
    print("[✓] HOÀN THÀNH! Tài liệu đã được tiền xử lý và lưu vào database")
    print("=" * 70)

# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """Main entry point"""
    if len(sys.argv) > 1:
        # File path provided as command-line argument
        file_path = sys.argv[1]
    else:
        # Prompt user for file path
        print("Nhập đường dẫn đến file dữ liệu (PDF hoặc DOCX):")
        file_path = input().strip()
    
    if not file_path:
        print("[!] Lỗi: Không nhập đường dẫn file")
        sys.exit(1)
    
    unified_pipeline(file_path)

if __name__ == "__main__":
    main()
